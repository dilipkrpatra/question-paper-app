import os
import random
from io import BytesIO
from datetime import datetime

import pandas as pd
import streamlit as st

import streamlit.components.v1 as components

from docx import Document

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer
)

from reportlab.lib.styles import (
    getSampleStyleSheet,
    ParagraphStyle
)

from reportlab.lib.enums import TA_CENTER

from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


# ---------------------------------------------------
# PAGE CONFIG
# ---------------------------------------------------

st.set_page_config(
    page_title="Question Paper Generator",
    page_icon="📘",
    layout="wide"
)

# ---------------------------------------------------
# FONTS
# ---------------------------------------------------

pdfmetrics.registerFont(
    TTFont(
        "Bengali",
        "fonts/NotoSansBengali-Regular.ttf"
    )
)

pdfmetrics.registerFont(
    TTFont(
        "Bengali2",
        "fonts/solaimanlipi_22-02-2012.ttf"
    )
)


# ---------------------------------------------------
# SESSION STATE
# ---------------------------------------------------

defaults = {
    "generated": {},
    "qa_text": "",
    "ans_text": "",
    "qa_ans_text": ""
}

for key, value in defaults.items():
    if key not in st.session_state:
        st.session_state[key] = value


# ---------------------------------------------------
# CACHE WORKBOOK
# ---------------------------------------------------

@st.cache_data(show_spinner=False)
def load_question_bank(file_path):

    workbook = pd.read_excel(
        file_path,
        sheet_name=None,
        header=None,
        engine="openpyxl"
    )

    bank = {}

    for sheet, df in workbook.items():

        heading = str(df.iloc[0, 0])

        # ONLY take first 2 columns (ignore notes/extra columns safely)
        df = df.iloc[1:, [0, 1]]

        # Rename for safety
        df.columns = ["question", "answer"]

        # Drop rows where question or answer is missing ONLY in these 2 columns
        df = df.dropna(subset=["question", "answer"])

        qa_pairs = list(zip(
            df["question"].astype(str),
            df["answer"].astype(str)
        ))

        bank[sheet] = {
            "heading": heading,
            "questions": qa_pairs
        }

    return bank

# ---------------------------------------------------
# PDF CREATOR
# ---------------------------------------------------

def create_pdf(title, text, subject):

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer)
    styles = getSampleStyleSheet()

    header_style = ParagraphStyle(
        "Header",
        parent=styles["Heading1"],
        alignment=TA_CENTER,
        fontSize=18,
        spaceAfter=12
    )

    bengali_style = ParagraphStyle(
        "Bengali",
        parent=styles["Normal"],
        fontName="Bengali2",
        fontSize=11,
        leading=16
    )

    story = []

    story.append(
        Paragraph(title, header_style)
    )

    story.append(
        Paragraph(
            f"Subject : {subject}",
            styles["Normal"]
        )
    )

    story.append(
        Paragraph(
            f"Date : {datetime.now().strftime('%d-%m-%Y')}",
            styles["Normal"]
        )
    )

    story.append(
        Spacer(1, 12)
    )

    for line in text.split("\n"):

        if line.strip():
            story.append(
                Paragraph(
                    line,
                    bengali_style
                )
            )

            story.append(
                Spacer(1, 5)
            )

        else:
            story.append(
                Spacer(1, 10)
            )

    doc.build(story)
    buffer.seek(0)

    return buffer


# ---------------------------------------------------
# DOCX CREATOR
# ---------------------------------------------------

def create_docx(title, text):

    doc = Document()
    doc.add_heading(title, level=1)

    for line in text.split("\n"):

        doc.add_paragraph(line)

    buffer = BytesIO()

    doc.save(buffer)

    buffer.seek(0)

    return buffer


# ---------------------------------------------------
# BUILD QUESTION TEXT
# ---------------------------------------------------

def build_text(selected_questions, bank):

    qa_text = ""

    ans_text = ""

    qa_ans_text = ""

    for sheet, questions in selected_questions.items():

        heading = bank[sheet]["heading"]
        qa_text += f"\n{heading}\n"
        qa_text += "-" * 40 + "\n"
        ans_text += f"\n{heading}\n"
        ans_text += "-" * 40 + "\n"
        qa_ans_text += f"\n{heading}\n"
        qa_ans_text += "-" * 40 + "\n"

        for i, (q, a) in enumerate(
            questions,
            1
        ):

            qa_text += f"{i}. {q}\n"
            ans_text += f"{i}. {a}\n"
            qa_ans_text += f"{i}. {q}  ⮫  {a}\n"

    return qa_text, ans_text, qa_ans_text

# =====================================================
# MAIN TITLE
# =====================================================

st.title("📘 Test Paper Generator")


# =====================================================
# SUBJECT SELECTION
# =====================================================

QUESTION_BANK_FOLDER = "question_bank"

files = sorted([
    f for f in os.listdir(QUESTION_BANK_FOLDER)
    if f.endswith((".xlsx", ".xls"))
])

if not files:
    st.error("No Excel files found in 'question_bank' folder.")
    st.stop()

selected_file = st.selectbox(
    "📂 Select Subject",
    files
)

file_path = os.path.join(
    QUESTION_BANK_FOLDER,
    selected_file
)

subject_name = os.path.splitext(selected_file)[0]


# =====================================================
# LOAD QUESTION BANK (ONLY ONCE)
# =====================================================

with st.spinner("Loading Question Bank..."):

    bank = load_question_bank(file_path)


sheet_names = list(bank.keys())


# =====================================================
# TOPIC SELECTION
# =====================================================

st.subheader("📚 Select Topics")

topic_display = {}

for sheet in sheet_names:

    total = len(bank[sheet]["questions"])

    topic_display[
        f"{sheet} ({total})"
    ] = sheet


selected_display = st.pills(
    "Choose Topics",
    list(topic_display.keys()),
    selection_mode="multi"
)

selected_topics = [
    topic_display[item]
    for item in selected_display
]


# =====================================================
# QUESTION COUNT
# =====================================================

topic_counts = {}

if selected_topics:

    st.subheader("Number of Questions")

    cols = st.columns(3)

    for index, sheet in enumerate(selected_topics):

        with cols[index % 3]:

            maximum = len(bank[sheet]["questions"])

            topic_counts[sheet] = st.number_input(

                sheet,
                min_value=1,
                max_value=maximum,
                value=min(5, maximum),
                step=1,
                key=f"count_{sheet}"
            )


# =====================================================
# BUTTON CSS
# =====================================================

st.markdown("""

<style>

div.stButton > button:first-child{

    background:linear-gradient(90deg,#1e3c72,#2a5298);
    color:white;
    font-size:22px;
    font-weight:bold;
    height:58px;
    border:none;
    border-radius:12px;

}

div.stButton > button:first-child:hover{

    background:linear-gradient(90deg,#2a5298,#1e3c72);

}

</style>

""", unsafe_allow_html=True)


# =====================================================
# GENERATE BUTTON
# =====================================================

if selected_topics:

    generate = st.button(
        "⚡ GENERATE QUESTION PAPER",
        use_container_width=True
    )

else:
    generate = False

# =====================================================
# GENERATE RANDOM QUESTIONS
# =====================================================

if generate:

    generated = {}
    for sheet in selected_topics:
        questions = bank[sheet]["questions"]
        count = min(

            topic_counts[sheet],
            len(questions)
        )

        generated[sheet] = random.sample(
            questions,
            count
        )

    st.session_state.generated = generated

    qa_text, ans_text, qa_ans_text = build_text(
        generated,
        bank
    )

    st.session_state.qa_text = qa_text
    st.session_state.ans_text = ans_text
    st.session_state.qa_ans_text = qa_ans_text

    st.success("✅ Question Paper Generated Successfully")
    
# =====================================================
# PREVIEW
# =====================================================

st.divider()

st.subheader("📄 Question Paper")

if st.session_state.qa_text == "":

    st.info("Generate a question paper to see the preview.")

else:

    st.code(
        st.session_state.qa_text,
        language="text"
    )

# =====================================================
# ANSWER KEY
# =====================================================

if st.session_state.qa_text != "":

    show_answers = st.toggle(
        "👁 Show Answers",
        value=False
    )

    if show_answers:

        st.subheader("📘 Answer Key")
        st.code(
            st.session_state.qa_ans_text,
            language="text"
        )

